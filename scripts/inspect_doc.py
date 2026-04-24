"""
Inspect the Word document structure to understand how it's formatted.
Run: python scripts/inspect_doc.py data/dictionary.docx
"""
import sys
from docx import Document

path = sys.argv[1] if len(sys.argv) > 1 else 'data/dictionary.docx'
doc = Document(path)

print("=== TABLES ===")
print(f"Number of tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables[:2]):
    print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    for row in table.rows[:5]:
        print("  ROW:")
        for j, cell in enumerate(row.cells):
            text = cell.text.strip()[:120]
            print(f"    Cell {j}: {repr(text)}")

print("\n=== FIRST 30 PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs[:30]):
    text = ''.join(run.text for run in para.runs).strip()
    if text:
        print(f"  [{i}] style={repr(para.style.name)} text={repr(text[:120])}")
